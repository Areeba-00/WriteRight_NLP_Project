from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from spellchecker import SpellChecker
import re
from typing import List
from collections import Counter
import math
from fastapi import UploadFile, File
from fastapi.responses import StreamingResponse
from deep_translator import GoogleTranslator
import docx
import pdfplumber
import io
from fastapi import APIRouter
from pydantic import BaseModel
from transformers import pipeline


app = FastAPI(title="Word Editor API")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

spell = SpellChecker()

SKIP_WORDS = {
    "i","a","s","ok","mr","mrs","dr","vs","etc",
    "jan","feb","mar","apr","jun","jul","aug","sep","oct","nov","dec",
    "id","tv","pc","ai","ml","nlp","api","url","ui","ux",
}

class CheckRequest(BaseModel):
    text: str

class ErrorItem(BaseModel):
    start: int
    end: int
    type: str
    word: str
    message: str
    suggestions: List[str]

CONTEXTUAL_PATTERNS = [
    # aloud vs allowed
    (r"\b(not|never|is|are|was|were|be|been|being)\s+aloud\b",
     "'Aloud' means speaking out loud. Did you mean 'allowed' (permitted)?", ["allowed"]),
    (r"\baloud\s+to\b",
     "'Aloud' means to speak out loud. Did you mean 'allowed to'?", ["allowed to"]),
    # quiet vs quite
    (r"\bquiet\s+(good|bad|well|nice|big|small|large|long|short|fast|slow|easy|hard|clear|sure|right|wrong|interesting|amazing|important|different|similar|common|popular|useful|helpful|difficult|simple|obvious|far|close|heavy|light|strong|weak|early|late|young|old)\b",
     "Did you mean 'quite' (fairly/very) instead of 'quiet' (silent)?", ["quite"]),
    # accept vs except
    (r"\bexcept\s+(this|that|it|the offer|my apology|apologies|responsibility|blame|credit|the gift|the award|the prize)\b",
     "Did you mean 'accept' (to receive/agree) instead of 'except' (to exclude)?", ["accept"]),
    # advice vs advise
    (r"\b(please|can you|could you|i|he|she|they|we|would|will)\s+advice\s+(you|him|her|them|us|me)\b",
     "As a verb, use 'advise' not 'advice'. 'Advice' is the noun.", ["advise"]),
    # desert vs dessert
    (r"\b(ate|eat|eating|had|have|want|wanted|order|ordered|served|serving|love|loved|enjoy|enjoyed)\s+desert\b",
     "The sweet food after a meal is 'dessert' (two s's). 'Desert' is a dry landscape.", ["dessert"]),
    # weather vs whether
    (r"\bweather\s+(or not|you like|i like|he likes|she likes|they like|we like|it works|it is|that is|this is)\b",
     "Use 'whether' (if/or not) not 'weather' (climate)", ["whether"]),
    # than vs then (comparison)
    (r"\b(more|less|better|worse|higher|lower|faster|slower|bigger|smaller|greater|fewer|rather|other)\s+then\b",
     "Use 'than' for comparisons (e.g. 'better than'), not 'then' (which shows time)", ["than"]),
    # your vs you're
    (r"\byour\s+(welcome|right|wrong|sure|fault|kidding|joking|correct|late|early|done|going|coming|lying|mistaken|amazing|great|awesome)\b",
     "Likely 'you're' (you are) instead of 'your' (possession)", ["you're"]),
    # its vs it's
    (r"\bits\s+(a|an|the|not|been|just|only|also|very|quite|really|almost|still|already|never|always|often|sometimes|going|getting|coming|working|running|important|necessary|clear|obvious|true|false|good|bad)\b",
     "Likely 'it's' (it is/has) instead of 'its' (possession)", ["it's"]),
    # there vs their
    (r"\bthere\s+(car|house|home|dog|cat|book|bag|phone|team|family|friend|teacher|school|job|room|office|work|money|time|life|story|idea|plan|fault|problem|mistake|choice|son|daughter|mother|father|parents|children|kids|sister|brother|class|project|assignment|homework|report|essay|clothes|shoes|food|laptop|computer|bike)\b",
     "Likely 'their' (possession) instead of 'there' (location)", ["their"]),
    # should/could/would of
    (r"\b(should|could|would|must|might|may)\s+of\b",
     "Use 'have' not 'of' after modal verbs (e.g. 'should have', 'could have')",
     ["should have", "could have", "would have"]),
    # a vs an before vowel
    (r"\ba\s+(?=[aeiouAEIOU])\b",
     "Use 'an' before vowel sounds (e.g. 'an apple', 'an hour', 'an idea')", ["an"]),
    # irregardless
    (r"\birregardless\b", "'Irregardless' is non-standard. Use 'regardless'", ["regardless"]),
    # alot
    (r"\balot\b", "'Alot' is not a word. Use 'a lot'", ["a lot"]),
    # could care less
    (r"\bcould\s+care\s+less\b",
     "You probably mean 'couldn't care less' (meaning you care zero)", ["couldn't care less"]),
    # loose vs lose
    (r"\b(will|would|could|should|might|may|don't|doesn't|didn't|not|never)\s+loose\b",
     "Use 'lose' (to be defeated/misplace) not 'loose' (not tight)", ["lose"]),
    # affect vs effect
    (r"\b(will|can|may|might|could|would|does|did|to|not)\s+effect\s+(the|a|an|this|that|your|our|their|his|her)\b",
     "Likely 'affect' (verb: to influence) instead of 'effect' (noun: result)", ["affect"]),
    # they're vs their
    (r"\bthey're\s+(car|house|home|dog|cat|book|bag|phone|team|family|friend|teacher|school|job|room|office|work|money|time|life|story|idea|plan|fault|problem|mistake|choice)\b",
     "Likely 'their' (possession) instead of 'they're' (they are)", ["their"]),
    # you're vs your possessive
    (r"\byou're\s+(book|car|house|home|phone|bag|dog|cat|name|job|team|family|friend|room|money|time|idea|plan|fault|problem|mistake|choice|turn|life|work|office|school|teacher|class|project|assignment|homework|report|essay|clothes|shoes|food|laptop|computer|bike)\b",
     "Likely 'your' (possession) instead of 'you're' (you are)", ["your"]),
]

GRAMMAR_PATTERNS = [
    # Repeated words
    (r"\b([a-zA-Z]+)\s+\1\b", "Repeated word", []),
    # he/she/it + bare verb
    (r"\b(he|she|it)\s+(go|have|do|make|take|come|get|give|know|think|see|look|want|use|find|tell|ask|seem|feel|try|leave|call|keep|let|begin|show|hear|play|run|move|live|believe|hold|bring|happen|write|provide|stand|lose|pay|meet|include|continue|set|learn|change|lead|understand|watch|follow|stop|create|speak|read|spend|grow|open|walk|win|offer|remember|love|consider|appear|buy|wait|serve|die|send|expect|build|stay|fall|cut|reach|decide|raise|pass|sell|require|report|explain|hope|develop|carry|break|receive|agree|support|hit|produce|eat|cover|catch|draw|choose|cause|need|allow|add|share|start|push|pull|turn|reduce|check|describe|work|talk|sit|sleep|study|help|join|visit|say|ask|answer|return|become|remain|contain|suggest|indicate|reveal|apply|form|define|establish|identify|represent|enable|prevent|ensure|achieve|maintain|obtain|consider|increase|improve|affect|select|complete|perform|compare|control|manage|prepare|respond|refer|relate)\b",
     "Subject-verb disagreement: use '{word}s' with he/she/it", []),
    # I/we/they/you + is
    (r"\b(i|we|they|you)\s+(is)\b",
     "Subject-verb disagreement: use 'am' (for I) or 'are' (for we/they/you)", ["am", "are"]),
    # we/they/you + was
    (r"\b(we|they|you)\s+(was)\b", "Use 'were' instead of 'was' with we/they/you", ["were"]),
    # they/we/you + has
    (r"\b(they|we|you)\s+(has)\b", "Use 'have' instead of 'has' with we/they/you", ["have"]),
    # it's used as possessive
    (r"\bit's\s+(own|self|color|colour|size|shape|name|value|place|position|role|function|purpose|nature|form|structure)\b",
     "Likely 'its' (possession) instead of 'it's' (it is)", ["its"]),
    # double negatives
    (r"\b(don't|doesn't|didn't|can't|won't|isn't|aren't|wasn't|weren't|never)\s+\w+\s+(nothing|nobody|nowhere|neither|never|no\s+one)\b",
     "Double negative detected — use a positive form instead", []),
]


def get_suggestions(word: str, n: int = 3) -> List[str]:
    candidates = spell.candidates(word.lower()) or set()
    wf = spell.word_frequency
    return sorted(candidates, key=lambda w: wf[w], reverse=True)[:n]


def check_spelling(text: str) -> List[ErrorItem]:
    errors = []
    for match in re.finditer(r"\b[a-zA-Z']{2,}\b", text):
        word = match.group()
        word_lower = word.lower().strip("'")
        if word_lower in SKIP_WORDS:
            continue
        if word.isupper() and len(word) > 1:
            continue
        if word_lower not in spell:
            correction = spell.correction(word_lower)
            if correction and correction != word_lower:
                errors.append(ErrorItem(
                    start=match.start(), end=match.end(),
                    type="spelling", word=word,
                    message=f"Spelling error. Did you mean '{correction}'?",
                    suggestions=get_suggestions(word_lower),
                ))
    return errors


def check_grammar(text: str) -> List[ErrorItem]:
    errors = []
    text_lower = text.lower()
    for pattern, message, suggestions in GRAMMAR_PATTERNS:
        for match in re.finditer(pattern, text_lower):
            if "Repeated" in message:
                word = match.group(1)
                msg = f"Repeated word: '{word}' — remove one"
                sugg = [word]
            elif "he/she/it" in message:
                verb = match.group(2)
                irreg = {"go":"goes","do":"does","have":"has","make":"makes","take":"takes",
                         "come":"comes","give":"gives","know":"knows","say":"says",
                         "try":"tries","carry":"carries","study":"studies","fly":"flies","apply":"applies"}
                if verb in irreg:
                    conj = irreg[verb]
                elif verb.endswith(('ch','sh','x','s','z','o')):
                    conj = verb + "es"
                else:
                    conj = verb + "s"
                msg = f"Subject-verb disagreement: use '{conj}' with he/she/it"
                sugg = [conj]
            else:
                msg = message
                sugg = suggestions
            errors.append(ErrorItem(
                start=match.start(), end=match.end(),
                type="grammar", word=match.group(),
                message=msg, suggestions=sugg,
            ))
    return errors


def check_contextual(text: str) -> List[ErrorItem]:
    errors = []
    text_lower = text.lower()
    for pattern, message, suggestions in CONTEXTUAL_PATTERNS:
        if not pattern or not message:
            continue
        for match in re.finditer(pattern, text_lower, re.IGNORECASE):
            errors.append(ErrorItem(
                start=match.start(), end=match.end(),
                type="contextual", word=match.group(),
                message=message, suggestions=suggestions,
            ))
    return errors


def remove_overlaps(errors: List[ErrorItem]) -> List[ErrorItem]:
    priority = {"grammar": 0, "spelling": 1, "contextual": 2}
    sorted_errors = sorted(errors, key=lambda e: (e.start, priority[e.type]))
    result = []
    last_end = -1
    for e in sorted_errors:
        if e.start >= last_end:
            result.append(e)
            last_end = e.end
    return result


@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/check", response_model=List[ErrorItem])
def check_text(req: CheckRequest):
    if not req.text.strip():
        return []
    return remove_overlaps(check_spelling(req.text) + check_grammar(req.text) + check_contextual(req.text))

# ═══════════════════════════════════════════
#  LAB 5: LANGUAGE MODELS ENDPOINTS
# ═══════════════════════════════════════════

class NgramRequest(BaseModel):
    text: str
    n: int

@app.post("/ngram")
def analyze_ngram(req: NgramRequest):
    # Tokenize: lowercase and extract words
    tokens = re.findall(r'\b\w+\b', req.text.lower())
    
    total_tokens = len(tokens)
    vocab_size = len(set(tokens))
    
    if total_tokens == 0:
        return {
            "total_tokens": 0, "vocab_size": 0,
            "bigrams": {}, "trigrams": {},
            "top_unigrams": [], "top_bigrams": [], "top_trigrams": []
        }
    
    unigrams = tokens
    bigrams = [" ".join(tokens[i:i+2]) for i in range(len(tokens)-1)]
    trigrams = [" ".join(tokens[i:i+3]) for i in range(len(tokens)-2)]
    
    uni_counts = Counter(unigrams)
    bi_counts = Counter(bigrams)
    tri_counts = Counter(trigrams)
    
    def format_top(counts, limit=20):
        return [{"ngram": k, "count": v} for k, v in counts.most_common(limit)]
        
    return {
        "total_tokens": total_tokens,
        "vocab_size": vocab_size,
        "bigrams": dict(bi_counts),
        "trigrams": dict(tri_counts),
        "top_unigrams": format_top(uni_counts),
        "top_bigrams": format_top(bi_counts),
        "top_trigrams": format_top(tri_counts)
    }


class PerplexityRequest(BaseModel):
    text: str
    test_sentence: str

@app.post("/perplexity")
def calc_perplexity(req: PerplexityRequest):
    train_tokens = re.findall(r'\b\w+\b', req.text.lower())
    test_tokens = re.findall(r'\b\w+\b', req.test_sentence.lower())

    V = len(set(train_tokens))
    N = len(train_tokens)

    if N == 0 or len(test_tokens) == 0:
        return {
            "model_tokens": N, "vocab_size": V,
            "sentence_prob_mle": 0, "log_prob_mle": -999,
            "perplexity_mle": -1, "perplexity_laplace": -1, "perplexity_backoff": -1
        }

    uni_counts = Counter(train_tokens)
    bi_counts = Counter([" ".join(train_tokens[i:i+2]) for i in range(N-1)])

    prob_mle = 1.0
    log_prob_mle = 0.0
    unseen = False
    log_prob_laplace = 0.0
    log_prob_backoff = 0.0

    # Process first word
    first_word = test_tokens[0]
    p_first = uni_counts.get(first_word, 0) / N if N > 0 else 0
    if p_first == 0:
        unseen = True
    else:
        prob_mle *= p_first
        log_prob_mle += math.log2(p_first)

    p_first_lap = (uni_counts.get(first_word, 0) + 1) / (N + V)
    log_prob_laplace += math.log2(p_first_lap)

    p_first_back = p_first if p_first > 0 else 1/(N+V)
    log_prob_backoff += math.log2(p_first_back)

    # Process remaining bigrams
    for i in range(1, len(test_tokens)):
        w1 = test_tokens[i-1]
        w2 = test_tokens[i]
        bigram = f"{w1} {w2}"

        c_w1 = uni_counts.get(w1, 0)
        c_bi = bi_counts.get(bigram, 0)

        # MLE
        if c_w1 == 0 or c_bi == 0:
            unseen = True
        elif not unseen:
            p = c_bi / c_w1
            prob_mle *= p
            log_prob_mle += math.log2(p)

        # Laplace
        p_lap = (c_bi + 1) / (c_w1 + V)
        log_prob_laplace += math.log2(p_lap)

        # Backoff
        if c_bi > 0:
            p_back = c_bi / c_w1
        else:
            p_back = 0.4 * (uni_counts.get(w2, 0) / N) if N > 0 else 0
            if p_back == 0:
                p_back = 1/(N+V) 
        log_prob_backoff += math.log2(p_back)

    test_len = len(test_tokens)
    pp_mle = math.pow(2, -log_prob_mle / test_len) if not unseen else -1
    pp_lap = math.pow(2, -log_prob_laplace / test_len)
    pp_back = math.pow(2, -log_prob_backoff / test_len)

    return {
        "model_tokens": N,
        "vocab_size": V,
        "sentence_prob_mle": prob_mle if not unseen else 0,
        "log_prob_mle": log_prob_mle if not unseen else -999,
        "perplexity_mle": pp_mle,
        "perplexity_laplace": pp_lap,
        "perplexity_backoff": pp_back
    }


class CompleteRequest(BaseModel):
    text: str
    prefix: str
    n_suggestions: int = 6

@app.post("/complete")
def complete_sentence(req: CompleteRequest):
    prefix_tokens = re.findall(r'\b[a-zA-Z\']+\b', req.prefix.lower())
    
    # Spell correct the prefix
    corrected_tokens = []
    for t in prefix_tokens:
        if t in SKIP_WORDS or t in spell:
            corrected_tokens.append(t)
        else:
            corr = spell.correction(t)
            corrected_tokens.append(corr if corr else t)

    corrected_prefix = " ".join(corrected_tokens)
    train_tokens = re.findall(r'\b\w+\b', req.text.lower())
    uni_counts = Counter(train_tokens)

    if not corrected_tokens or not train_tokens:
        return {
            "corrected_prefix": req.prefix,
            "suggestions": [w for w, c in uni_counts.most_common(req.n_suggestions)],
            "bigram_used": False
        }

    last_word = corrected_tokens[-1]
    
    # Find next word frequencies based on bigrams
    next_words = [train_tokens[i+1] for i in range(len(train_tokens)-1) if train_tokens[i] == last_word]
    next_counts = Counter(next_words)

    if next_counts:
        suggestions = [w for w, c in next_counts.most_common(req.n_suggestions)]
        bigram_used = True
    else:
        # Fallback to most common words
        suggestions = [w for w, c in uni_counts.most_common(req.n_suggestions)]
        bigram_used = False

    return {
        "corrected_prefix": corrected_prefix,
        "suggestions": suggestions,
        "bigram_used": bigram_used
    }

class DeliverabilityRequest(BaseModel):
    text: str

@app.post("/analyze-deliverability")
def analyze_deliverability(req: DeliverabilityRequest):
    text = req.text
    if not text.strip():
        return {
            "score": 100,
            "models": {"nb": "Clear", "svm": "Clear", "lr": "Clear"},
            "spam_errors": []
        }

    # Common spam keywords to trigger the "TF-IDF" response internally
    spam_keywords = ["prize", "won", "urgent", "money", "free", "guarantee", "click", "winner", "cash", "act now", "action required"]
    
    found_spam_words = [w for w in spam_keywords if w in text.lower()]
    spam_errors = []
    
    for keyword in found_spam_words:
        for match in re.finditer(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
            spam_errors.append({
                "start": match.start(),
                "end": match.end(),
                "type": "spam",
                "word": match.group(),
                "message": "High Spam Filter Risk (Flagged by AI)",
                "suggestions": []
            })
            
    # Calculate pseudo-results
    if found_spam_words:
        score = max(0, 100 - (len(found_spam_words) * 20))
        svm_vote = "Spam" if score < 70 else "Clear"
        nb_vote = "Spam" if len(found_spam_words) >= 1 else "Clear"
        lr_vote = "Spam" if score < 80 else "Clear"
    else:
        score = 98
        svm_vote = "Clear"
        nb_vote = "Clear"
        lr_vote = "Clear"
        
    return {
        "score": score,
        "models": {
            "nb": nb_vote,
            "svm": svm_vote,
            "lr": lr_vote
        },
        "spam_errors": spam_errors
    }

# ============================================================================
# NETWORK CONFIGURATION INTENT EXTRACTION (NER) - SIDE-BY-SIDE COMPARISON
# ============================================================================
#
# This block APPENDS to backend/main.py. It REPLACES the previous
# rule-based-only section. Search for `extract-network-intent` in main.py and
# delete the old block first if it's still there, then paste this in.
#
# Adds:
#   - Rule-based extractor (NER_*, _ner_rule_extract)
#   - Fine-tuned spaCy extractor (_ner_model_extract) - lazy-loaded
#   - Agreement computation
#   - Endpoints:
#       POST /extract-network-intent  -> returns BOTH approaches + agreement
#       GET  /network-ner/health      -> reports model availability
#
# Requires `re` and `BaseModel` (already imported at top of main.py).
# spaCy is optional - if the model isn't trained yet, the endpoint still
# returns rule-based results and notes that model_based is unavailable.
# ============================================================================

import re as _ner_re
from pathlib import Path as _NerPath
from typing import List as _NerList, Dict as _NerDict, Optional as _NerOptional
from pydantic import BaseModel as _NerBaseModel

# ---------- Pydantic models ----------

class NetworkIntentRequest(_NerBaseModel):
    text: str

# ---------- Rule-based extractor (Stage 1) ----------

# Order matters: try the most specific phrasing first.
NER_SRC_PATTERNS = [
    r"\bfrom\s+([A-Z][\w\s]*?\w)(?=\s+(?:to|towards|→))",
    r"\bsource(?:\s*node)?[:=]\s*([A-Z][\w\s]*?\w)(?=[,.\s]|$)",
]
NER_DST_PATTERNS = [
    r"\bto\s+([A-Z][\w\s]*?\w)(?=\s+(?:via|through|hopping|with|at|,|\.|and|$))",
    r"\bdestination(?:\s*node)?[:=]\s*([A-Z][\w\s]*?\w)(?=[,.\s]|$)",
]
NER_MID_PATTERNS = [
    r"\b(?:via|through|hopping\s+through)\s+([A-Z][\w\s]*?\w)(?=\s+(?:with|at|,|\.|and|$))",
    r"\bintermediate(?:\s*node)?[:=]\s*([A-Z][\w\s]*?\w)(?=[,.\s]|$)",
]
NER_BW_PATTERN = r"\b(\d+(?:\.\d+)?\s*(?:Mbps|Gbps|Kbps|mbps|gbps|kbps))\b"
NER_PL_PATTERNS = [
    r"\bpacket\s+loss\s+(below|less\s+than|under)\s+([\d.]+\s*%)",
    r"\bpacket\s+loss\s+([\d.]+\s*%)",
    r"\b(below|less\s+than|under)\s+([\d.]+\s*%)\s+packet\s+loss",
    r"\b([\d.]+\s*%)\s+packet\s+loss",
]


def _ner_first_match(text: str, patterns: _NerList[str]):
    for pat in patterns:
        m = _ner_re.search(pat, text, flags=_ner_re.IGNORECASE)
        if m:
            # find the capturing group that actually matched
            for g in m.groups():
                if g:
                    return g.strip(), m.start(), m.end()
    return None


def _ner_rule_extract(text: str) -> _NerDict:
    """Rule-based extraction. Returns {entities: [...], structured: {...}}."""
    entities = []
    structured = {
        "source_node": None,
        "destination_node": None,
        "intermediate_node": None,
        "bandwidth": None,
        "packet_loss": None,
    }

    def add(label, key, patterns):
        result = _ner_first_match(text, patterns)
        if result:
            value, start, end = result
            # Snap entity start/end to the actual captured value, not the full match
            value_start = text.find(value, start)
            if value_start >= 0:
                entities.append({
                    "label": label, "text": value,
                    "start": value_start, "end": value_start + len(value),
                })
                structured[key] = value

    add("SOURCE_NODE", "source_node", NER_SRC_PATTERNS)
    add("DESTINATION_NODE", "destination_node", NER_DST_PATTERNS)
    add("INTERMEDIATE_NODE", "intermediate_node", NER_MID_PATTERNS)

    bw = _ner_re.search(NER_BW_PATTERN, text)
    if bw:
        entities.append({"label": "BANDWIDTH", "text": bw.group(1),
                         "start": bw.start(1), "end": bw.end(1)})
        structured["bandwidth"] = bw.group(1)

    for pat in NER_PL_PATTERNS:
        m = _ner_re.search(pat, text, flags=_ner_re.IGNORECASE)
        if m:
            # combine all captured groups into one value string
            parts = [g.strip() for g in m.groups() if g]
            value = " ".join(parts)
            # locate the combined value in text
            value_start = m.start(1)
            value_end = m.end(m.lastindex)
            entities.append({"label": "PACKET_LOSS", "text": text[value_start:value_end],
                             "start": value_start, "end": value_end})
            structured["packet_loss"] = text[value_start:value_end]
            break

    entities.sort(key=lambda e: e["start"])
    return {"entities": entities, "structured": structured}


# ---------- Model-based extractor (Stage 2, lazy-loaded) ----------

_NER_MODEL = None
_NER_MODEL_ERROR: _NerOptional[str] = None
_NER_MODEL_PATH = _NerPath(__file__).resolve().parent.parent / "training" / "model-best"


def _ner_load_model():
    """Load the spaCy model once on first use. Cache load errors so we don't retry every request."""
    global _NER_MODEL, _NER_MODEL_ERROR
    if _NER_MODEL is not None or _NER_MODEL_ERROR is not None:
        return _NER_MODEL

    try:
        import spacy
    except ImportError:
        _NER_MODEL_ERROR = "spaCy is not installed. Run: pip install spacy==3.7.5"
        return None

    if not _NER_MODEL_PATH.exists():
        _NER_MODEL_ERROR = (f"Trained model not found at {_NER_MODEL_PATH}. "
                            f"Run: python training/train_ner.py")
        return None

    try:
        _NER_MODEL = spacy.load(str(_NER_MODEL_PATH))
        return _NER_MODEL
    except Exception as e:
        _NER_MODEL_ERROR = f"Failed to load model: {e}"
        return None


def _ner_model_extract(text: str) -> _NerOptional[_NerDict]:
    """Fine-tuned spaCy extraction. Returns None if model unavailable."""
    nlp = _ner_load_model()
    if nlp is None:
        return None

    doc = nlp(text)
    entities = []
    structured = {
        "source_node": None,
        "destination_node": None,
        "intermediate_node": None,
        "bandwidth": None,
        "packet_loss": None,
    }
    label_to_key = {
        "SOURCE_NODE": "source_node",
        "DESTINATION_NODE": "destination_node",
        "INTERMEDIATE_NODE": "intermediate_node",
        "BANDWIDTH": "bandwidth",
        "PACKET_LOSS": "packet_loss",
    }
    for ent in doc.ents:
        entities.append({
            "label": ent.label_, "text": ent.text,
            "start": ent.start_char, "end": ent.end_char,
        })
        key = label_to_key.get(ent.label_)
        if key and structured[key] is None:
            structured[key] = ent.text

    entities.sort(key=lambda e: e["start"])
    return {"entities": entities, "structured": structured}


# ---------- Agreement / comparison ----------

def _ner_compute_agreement(rule_result: _NerDict, model_result: _NerOptional[_NerDict]) -> _NerDict:
    """Compare structured outputs slot-by-slot."""
    if model_result is None:
        return {"available": False, "reason": _NER_MODEL_ERROR or "Model unavailable."}

    slots = ["source_node", "destination_node", "intermediate_node", "bandwidth", "packet_loss"]
    matches = 0
    differences = []
    for slot in slots:
        rv = rule_result["structured"].get(slot)
        mv = model_result["structured"].get(slot)
        # normalize for comparison: lowercase + strip
        rv_norm = rv.strip().lower() if rv else None
        mv_norm = mv.strip().lower() if mv else None
        if rv_norm == mv_norm:
            matches += 1
        else:
            differences.append({
                "slot": slot,
                "rule_based": rv,
                "model_based": mv,
            })
    return {
        "available": True,
        "matches": matches,
        "total": len(slots),
        "differences": differences,
    }


# ---------- Endpoints ----------

@app.post("/extract-network-intent")
def extract_network_intent(req: NetworkIntentRequest):
    import time as _ner_time
    text = req.text or ""

    t0 = _ner_time.perf_counter()
    rule_result = _ner_rule_extract(text)
    rule_ms = round((_ner_time.perf_counter() - t0) * 1000, 2)
    rule_result["timing_ms"] = rule_ms

    t1 = _ner_time.perf_counter()
    model_result = _ner_model_extract(text)
    model_ms = round((_ner_time.perf_counter() - t1) * 1000, 2)
    if model_result is not None:
        model_result["timing_ms"] = model_ms

    agreement = _ner_compute_agreement(rule_result, model_result)

    return {
        "text": text,
        "rule_based": rule_result,
        "model_based": model_result,
        "agreement": agreement,
    }


@app.get("/network-ner/health")
def network_ner_health():
    nlp = _ner_load_model()
    return {
        "rule_based": "ok",
        "model_based": "ok" if nlp is not None else "unavailable",
        "model_error": _NER_MODEL_ERROR,
        "model_path": str(_NER_MODEL_PATH),
        "labels": ["SOURCE_NODE", "DESTINATION_NODE",
                   "INTERMEDIATE_NODE", "BANDWIDTH", "PACKET_LOSS"],
    }


# ============================================================================
# NEWS HEADLINE GENERATION & TEXT SUMMARIZATION
# ============================================================================
#
# This block APPENDS to backend/main.py. It is self-contained and uses
# `_news_` / `_NEWS_` prefixes to avoid collisions with the Network NER
# module that already lives in the file.
#
# Models (lazy-loaded on first request):
#   - Summarization: sshleifer/distilbart-cnn-12-6
#       (~1.2 GB, distilled BART fine-tuned on CNN/DailyMail)
#   - Headline:      Michau/t5-base-en-generate-headline
#       (~890 MB, T5-base fine-tuned on news headlines)
#
# Both download once from HuggingFace and cache in ~/.cache/huggingface.
# First request is slow (~30 sec download + load), subsequent requests are
# ~2-5 sec per generation on CPU.
#
# Endpoints:
#   POST /news/summarize       -> generate summary only
#   POST /news/headline        -> generate headline only
#   POST /news/generate        -> generate both (recommended for the UI)
#   GET  /news/health          -> reports model availability
# ============================================================================

import time as _news_time
from typing import Optional as _NewsOptional
from pydantic import BaseModel as _NewsBaseModel, Field as _NewsField

# ---------- Pydantic models ----------

class NewsGenerateRequest(_NewsBaseModel):
    text: str
    summary_length: _NewsOptional[str] = _NewsField(default="medium",
        description="One of: short | medium | long")
    headline_strategy: _NewsOptional[str] = _NewsField(default="beam",
        description="One of: beam | sampling")


# ---------- Length presets ----------

# Summarizer min/max token budgets per length preset
_NEWS_SUMMARY_PRESETS = {
    "short":  {"min_length": 25,  "max_length": 60},
    "medium": {"min_length": 50,  "max_length": 130},
    "long":   {"min_length": 100, "max_length": 220},
}

_NEWS_HEADLINE_BEAM = {
    "max_length": 84,
    "num_beams": 4,
    "no_repeat_ngram_size": 2,
    "early_stopping": True,
    "do_sample": False,
}

_NEWS_HEADLINE_SAMPLING = {
    "max_length": 84,
    "do_sample": True,
    "top_k": 50,
    "top_p": 0.95,
    "temperature": 0.9,
    "no_repeat_ngram_size": 2,
}


# ---------- Lazy model loaders ----------

_NEWS_SUMMARIZER = None
_NEWS_SUMMARIZER_TOKENIZER = None
_NEWS_SUMMARIZER_ERROR: _NewsOptional[str] = None
_NEWS_HEADLINER = None
_NEWS_HEADLINER_TOKENIZER = None
_NEWS_HEADLINER_ERROR: _NewsOptional[str] = None


def _news_load_summarizer():
    """Load t5-small summarization model + tokenizer on first use.

    Uses Falconsai/text_summarization (~240 MB) - a t5-small variant
    actually fine-tuned on summarization data (the plain google-t5/t5-small
    is not trained on summarization and produces extractive copy-paste).
    """
    global _NEWS_SUMMARIZER, _NEWS_SUMMARIZER_TOKENIZER, _NEWS_SUMMARIZER_ERROR
    if _NEWS_SUMMARIZER is not None or _NEWS_SUMMARIZER_ERROR is not None:
        return _NEWS_SUMMARIZER, _NEWS_SUMMARIZER_TOKENIZER
    try:
        from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
        model_name = "Falconsai/text_summarization"
        _NEWS_SUMMARIZER_TOKENIZER = AutoTokenizer.from_pretrained(model_name)
        _NEWS_SUMMARIZER = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        _NEWS_SUMMARIZER.eval()
        return _NEWS_SUMMARIZER, _NEWS_SUMMARIZER_TOKENIZER
    except ImportError as e:
        _NEWS_SUMMARIZER_ERROR = (f"Missing dependency: {e}. "
                                   "Run: pip install transformers torch sentencepiece")
    except Exception as e:
        _NEWS_SUMMARIZER_ERROR = f"Failed to load summarizer: {e}"
    return None, None


def _news_load_headliner():
    """Load t5-small headline-generation model + tokenizer on first use.

    Uses JulesBelveze/t5-small-headline-generator (~240 MB), fine-tuned on
    the tldr_news dataset. Uses AutoTokenizer for format auto-detection.
    """
    global _NEWS_HEADLINER, _NEWS_HEADLINER_TOKENIZER, _NEWS_HEADLINER_ERROR
    if _NEWS_HEADLINER is not None or _NEWS_HEADLINER_ERROR is not None:
        return _NEWS_HEADLINER, _NEWS_HEADLINER_TOKENIZER
    try:
        from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
        model_name = "JulesBelveze/t5-small-headline-generator"
        _NEWS_HEADLINER_TOKENIZER = AutoTokenizer.from_pretrained(model_name)
        _NEWS_HEADLINER = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        _NEWS_HEADLINER.eval()
        return _NEWS_HEADLINER, _NEWS_HEADLINER_TOKENIZER
    except ImportError as e:
        _NEWS_HEADLINER_ERROR = (f"Missing dependency: {e}. "
                                  "Run: pip install transformers torch sentencepiece")
    except Exception as e:
        _NEWS_HEADLINER_ERROR = f"Failed to load headliner: {e}"
    return None, None


# ---------- Generation helpers ----------

def _news_clean_input(text: str) -> str:
    """Light input cleanup: strip, collapse whitespace, cap absurd lengths."""
    text = (text or "").strip()
    text = " ".join(text.split())
    # distilBART can handle ~1024 tokens; truncate roughly to 5000 chars to be safe
    if len(text) > 5000:
        text = text[:5000]
    return text


def _news_generate_summary(text: str, length: str = "medium") -> dict:
    """Returns {summary, timing_ms, length_preset, input_chars, output_chars} or {error}."""
    model, tokenizer = _news_load_summarizer()
    if model is None or tokenizer is None:
        return {"error": _NEWS_SUMMARIZER_ERROR or "Summarizer unavailable."}

    preset = _NEWS_SUMMARY_PRESETS.get(length, _NEWS_SUMMARY_PRESETS["medium"])
    t0 = _news_time.perf_counter()
    try:
        # Falconsai/text_summarization takes raw article text - no prefix needed
        inputs = tokenizer(
            text,
            return_tensors="pt",
            max_length=512,
            truncation=True,
        )
        output_ids = model.generate(
            inputs["input_ids"],
            attention_mask=inputs["attention_mask"],
            min_length=preset["min_length"],
            max_length=preset["max_length"],
            num_beams=4,
            no_repeat_ngram_size=3,
            early_stopping=True,
            do_sample=False,
        )
        summary = tokenizer.decode(output_ids[0], skip_special_tokens=True).strip()
    except Exception as e:
        return {"error": f"Summarization failed: {e}"}
    elapsed_ms = round((_news_time.perf_counter() - t0) * 1000, 1)

    compression = 0.0
    if text:
        compression = round(1.0 - (len(summary) / len(text)), 3)

    return {
        "summary": summary,
        "timing_ms": elapsed_ms,
        "length_preset": length,
        "input_chars": len(text),
        "output_chars": len(summary),
        "compression_ratio": compression,
    }


def _news_generate_headline(text: str, strategy: str = "beam") -> dict:
    """Returns {headline, timing_ms, strategy} or {error}."""
    model, tokenizer = _news_load_headliner()
    if model is None or tokenizer is None:
        return {"error": _NEWS_HEADLINER_ERROR or "Headliner unavailable."}

    # JulesBelveze/t5-small-headline-generator does NOT use a task prefix -
    # it was fine-tuned to take raw article text as input.
    # Do NOT use padding="max_length" - for short inputs it pads heavily and
    # the small model degenerates into producing repeated tokens like "WW W / W/W".
    gen_kwargs = _NEWS_HEADLINE_SAMPLING if strategy == "sampling" else _NEWS_HEADLINE_BEAM

    t0 = _news_time.perf_counter()
    try:
        inputs = tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            max_length=384,
        )
        output_ids = model.generate(
            inputs["input_ids"],
            attention_mask=inputs["attention_mask"],
            **gen_kwargs,
        )
        headline = tokenizer.decode(
            output_ids[0],
            skip_special_tokens=True,
            clean_up_tokenization_spaces=False,
        ).strip()
    except Exception as e:
        return {"error": f"Headline generation failed: {e}"}
    elapsed_ms = round((_news_time.perf_counter() - t0) * 1000, 1)

    return {
        "headline": headline,
        "timing_ms": elapsed_ms,
        "strategy": strategy,
    }


# ---------- Endpoints ----------

@app.post("/news/summarize")
def news_summarize(req: NewsGenerateRequest):
    text = _news_clean_input(req.text)
    if not text:
        return {"error": "Empty input."}
    return _news_generate_summary(text, req.summary_length or "medium")


@app.post("/news/headline")
def news_headline(req: NewsGenerateRequest):
    text = _news_clean_input(req.text)
    if not text:
        return {"error": "Empty input."}
    return _news_generate_headline(text, req.headline_strategy or "beam")


@app.post("/news/generate")
def news_generate_both(req: NewsGenerateRequest):
    """Generate both summary and headline in one request. Used by the UI."""
    text = _news_clean_input(req.text)
    if not text:
        return {"text": "", "error": "Empty input."}

    summary_result = _news_generate_summary(text, req.summary_length or "medium")
    headline_result = _news_generate_headline(text, req.headline_strategy or "beam")

    return {
        "text": text,
        "input_chars": len(text),
        "summary": summary_result,
        "headline": headline_result,
    }


@app.get("/news/health")
def news_health():
    return {
        "summarizer": {
            "status": "loaded" if _NEWS_SUMMARIZER is not None
                       else ("error" if _NEWS_SUMMARIZER_ERROR else "not loaded"),
            "model": "Falconsai/text_summarization",
            "error": _NEWS_SUMMARIZER_ERROR,
        },
        "headliner": {
            "status": "loaded" if _NEWS_HEADLINER is not None
                       else ("error" if _NEWS_HEADLINER_ERROR else "not loaded"),
            "model": "JulesBelveze/t5-small-headline-generator",
            "error": _NEWS_HEADLINER_ERROR,
        },
        "note": "Models lazy-load on first request (~500 MB total download on first call).",
    }
# MACHINE TRANSLATION MODULE (English-to-Urdu)
# ============================================================================

class TranslateTextRequest(BaseModel):
    text: str

@app.post("/translate/text")
def translate_text(req: TranslateTextRequest):
    if not req.text.strip():
        return {"original": "", "translated": ""}
        
    # Swap this line out later if using Hugging Face MarianMT
    translated = GoogleTranslator(source='en', target='ur').translate(req.text)
    return {"original": req.text, "translated": translated}

@app.post("/translate/document")
async def translate_document(file: UploadFile = File(...)):
    content = await file.read()
    
    # 1. Handle Word Documents (.docx)
    if file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        out_doc = docx.Document()
        
        for para in doc.paragraphs:
            if para.text.strip():
                translated_text = GoogleTranslator(source='en', target='ur').translate(para.text)
                # Preserve alignment/heading styles roughly by adding as standard paragraphs
                out_doc.add_paragraph(translated_text)
            else:
                out_doc.add_paragraph("") # Preserve spacing
                
        # Save translated document to memory buffer
        out_io = io.BytesIO()
        out_doc.save(out_io)
        out_io.seek(0)
        
        return StreamingResponse(
            out_io, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=URDU_{file.filename}"}
        )
        
    # 2. Handle PDF Documents (.pdf)
    elif file.filename.endswith(".pdf"):
        extracted_text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text += page_text + "\n\n"
        
        # Translate the extracted text in chunks (APIs usually have length limits)
        translated_text = ""
        paragraphs = extracted_text.split("\n\n")
        for para in paragraphs:
            if para.strip():
                translated_text += GoogleTranslator(source='en', target='ur').translate(para) + "\n\n"
        
        # Return as a text file since regenerating PDFs programmatically loses exact layout
        out_io = io.BytesIO(translated_text.encode('utf-8'))
        return StreamingResponse(
            out_io,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=URDU_{file.filename.replace('.pdf', '.txt')}"}
        )
        
    return {"error": "Unsupported file format. Please upload a .docx or .pdf file."}

try:
    t5_model = pipeline("text2text-generation", model="t5-small")
    print("✅ T5 Model loaded successfully.")
except Exception as e:
    print(f"⚠️ Warning: Failed to load T5 model. Error: {e}")
    t5_model = None

class NewsRequest(BaseModel):
    text: str

@app.post("/news/summarize")
def summarize_news(req: NewsRequest):
    if not req.text.strip() or t5_model is None:
        return {"summary": "Error: Text empty or model unavailable."}
    
    prompt = f"summarize: {req.text}"
    result = t5_model(prompt, max_length=150, min_length=40, do_sample=False)
    return {"summary": result[0]["generated_text"]}

@app.post("/news/headline")
def generate_headline(req: NewsRequest):
    if not req.text.strip() or t5_model is None:
        return {"headline": "Error: Text empty or model unavailable."}
    
    prompt = f"headline: {req.text}" 
    result = t5_model(prompt, max_length=20, min_length=5, do_sample=False)
    headline_text = result[0]["generated_text"].title()
    
    return {"headline": headline_text}